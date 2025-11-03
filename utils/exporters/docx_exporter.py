"""
DOCX Exporter - Export in formato Microsoft Word
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from html.parser import HTMLParser

from utils.exporters.base_exporter import BaseExporter
from utils.logger import AppLogger


class HTMLToDocxConverter(HTMLParser):
    """Convert HTML to DOCX runs with formatting"""

    def __init__(self, paragraph):
        super().__init__()
        self.paragraph = paragraph
        self.current_run = None
        self.format_stack = {
            'bold': False,
            'italic': False,
            'underline': False
        }
        self.ignore_content = False  # Flag to ignore content in <style>, <head>, etc.

    def handle_starttag(self, tag, attrs):
        # Ignore content in style and head tags
        if tag in ['style', 'head', 'script']:
            self.ignore_content = True
            return

        # Ignore structural tags
        if tag in ['html', 'body']:
            return

        if tag in ['b', 'strong']:
            self.format_stack['bold'] = True
        elif tag in ['i', 'em']:
            self.format_stack['italic'] = True
        elif tag == 'u':
            self.format_stack['underline'] = True
        elif tag == 'br':
            # Add line break
            if self.current_run:
                self.current_run.add_break()
        elif tag == 'span':
            # Handle Qt-style inline formatting with <span style="...">
            style_dict = dict(attrs)
            style_str = style_dict.get('style', '')

            # Check for bold (font-weight:700 or font-weight:bold)
            if 'font-weight:700' in style_str or 'font-weight:bold' in style_str:
                self.format_stack['bold'] = True

            # Check for italic
            if 'font-style:italic' in style_str:
                self.format_stack['italic'] = True

            # Check for underline
            if 'text-decoration: underline' in style_str:
                self.format_stack['underline'] = True

    def handle_endtag(self, tag):
        # Re-enable content after style/head tags
        if tag in ['style', 'head', 'script']:
            self.ignore_content = False
            return

        # Ignore structural tags
        if tag in ['html', 'body']:
            return

        if tag in ['b', 'strong']:
            self.format_stack['bold'] = False
        elif tag in ['i', 'em']:
            self.format_stack['italic'] = False
        elif tag == 'u':
            self.format_stack['underline'] = False
        elif tag == 'span':
            # Reset all formatting when closing span
            # (Qt typically uses one span per format)
            self.format_stack['bold'] = False
            self.format_stack['italic'] = False
            self.format_stack['underline'] = False

    def handle_data(self, data):
        # Ignore data if we're in a style/head tag
        if self.ignore_content:
            return

        # Always add data, even if it's just whitespace (spaces between words are important)
        if data:
            # Create run with current formatting
            self.current_run = self.paragraph.add_run(data)
            self.current_run.bold = self.format_stack['bold']
            self.current_run.italic = self.format_stack['italic']
            self.current_run.underline = self.format_stack['underline']


class DOCXExporter(BaseExporter):
    """
    Export DOCX compatibile con Microsoft Word.

    Supporta:
    - Stili nativi Word
    - Indice automatico
    - Formattazione preservata
    """

    def export(self, output_path: str) -> bool:
        """
        Export in DOCX

        Args:
            output_path: Percorso file DOCX destinazione

        Returns:
            bool: True se successo
        """
        try:
            AppLogger.info(f"DOCX Export: Starting export to {output_path}")

            # Get type-specific formatting
            self.type_formatting = self.get_type_specific_formatting()
            AppLogger.debug(f"Using type-specific formatting for {self.project.project_type}")

            # Crea documento Word
            doc = Document()

            # Imposta margini
            self._set_margins(doc)

            # Aggiungi cover page
            if self._get_option('include_cover', True):
                self._add_cover_page(doc)
                doc.add_page_break()

            # Aggiungi table of contents
            if self._get_option('include_toc', True):
                self._add_table_of_contents(doc)
                doc.add_page_break()

            # Aggiungi contenuto capitoli e scene
            self._add_manuscript_content(doc)

            # Salva documento
            doc.save(output_path)

            AppLogger.info(f"DOCX Export completed successfully: {output_path}")
            return True

        except Exception as e:
            AppLogger.error(f"DOCX Export failed: {e}", exc_info=True)
            return False

    def _set_margins(self, doc):
        """
        Imposta i margini del documento

        Args:
            doc: Documento Word
        """
        sections = doc.sections
        for section in sections:
            # Margini in inches (convertiti da cm se specificati nelle opzioni) - usa type-specific formatting
            fmt = self.type_formatting
            margin_left = self._get_option('margin_left', fmt.get('margin_left', 2.5))
            margin_right = self._get_option('margin_right', fmt.get('margin_right', 2.5))
            margin_top = self._get_option('margin_top', 2.5)
            margin_bottom = self._get_option('margin_bottom', 2.5)

            # Converti cm in inches (1 cm = 0.393701 inches)
            section.left_margin = Inches(margin_left * 0.393701)
            section.right_margin = Inches(margin_right * 0.393701)
            section.top_margin = Inches(margin_top * 0.393701)
            section.bottom_margin = Inches(margin_bottom * 0.393701)

    def _add_cover_page(self, doc):
        """
        Aggiunge la pagina di copertina

        Args:
            doc: Documento Word
        """
        # Spazio iniziale
        for _ in range(5):
            doc.add_paragraph()

        # Titolo
        title = doc.add_heading(self.project.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_paragraph(title, font_size=28, bold=True, color=RGBColor(44, 62, 80))

        # Spazio
        doc.add_paragraph()

        # Autore
        if self.project.author:
            author_para = doc.add_paragraph(f"by {self.project.author}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._format_paragraph(author_para, font_size=16, color=RGBColor(52, 73, 94))

        # Spazio
        for _ in range(2):
            doc.add_paragraph()

        # Info progetto
        if self.project.project_type:
            # Convert ProjectType enum to string
            project_type_str = self.project.project_type.get_display_name(self.project.language) if hasattr(self.project.project_type, 'get_display_name') else str(self.project.project_type.value)
            type_para = doc.add_paragraph(project_type_str)
            type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._format_paragraph(type_para, font_size=12, color=RGBColor(127, 140, 141))

        if self.project.genre:
            genre_para = doc.add_paragraph(self.project.genre)
            genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._format_paragraph(genre_para, font_size=12, color=RGBColor(127, 140, 141))

        # Spazio
        for _ in range(2):
            doc.add_paragraph()

        # Data
        export_date = datetime.now().strftime("%d %B %Y")
        date_para = doc.add_paragraph(export_date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_paragraph(date_para, font_size=12, color=RGBColor(127, 140, 141))

        AppLogger.debug("Added cover page to DOCX")

    def _add_table_of_contents(self, doc):
        """
        Aggiunge l'indice dei capitoli

        Args:
            doc: Documento Word
        """
        chapters = self._get_chapters()

        # Titolo TOC
        toc_title = doc.add_heading("Table of Contents", level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_paragraph(toc_title, font_size=22, bold=True, color=RGBColor(44, 62, 80))

        doc.add_paragraph()

        # Lista capitoli
        for i, chapter in enumerate(chapters, 1):
            chapter_title = chapter.title or f"Chapter {i}"
            scene_count = len(chapter.scenes)

            toc_text = f"Chapter {i}: {chapter_title}"
            if scene_count > 0:
                toc_text += f" ({scene_count} scene{'s' if scene_count != 1 else ''})"

            toc_para = doc.add_paragraph(toc_text)
            self._format_paragraph(toc_para, font_size=12, color=RGBColor(52, 73, 94))

        AppLogger.debug(f"Added table of contents with {len(chapters)} chapters")

    def _add_manuscript_content(self, doc):
        """
        Aggiunge il contenuto del manoscritto (capitoli e scene)

        Args:
            doc: Documento Word
        """
        chapters = self._get_chapters()
        chapter_font_size = self._get_option('chapter_font_size', 20)
        scene_font_size = self._get_option('scene_font_size', 14)
        content_font_size = self._get_option('content_font_size', 11)

        for i, chapter in enumerate(chapters, 1):
            # Titolo capitolo
            chapter_title = chapter.title or f"Chapter {i}"
            chapter_heading = doc.add_heading(f"Chapter {i}: {chapter_title}", level=1)
            self._format_paragraph(
                chapter_heading,
                font_size=chapter_font_size,
                bold=True,
                color=RGBColor(44, 62, 80)
            )

            # Scene del capitolo
            for j, scene in enumerate(chapter.scenes, 1):
                # Titolo scena
                scene_title = scene.title or f"Scene {j}"
                scene_heading = doc.add_heading(scene_title, level=2)
                self._format_paragraph(
                    scene_heading,
                    font_size=scene_font_size,
                    bold=True,
                    color=RGBColor(52, 73, 94)
                )

                # Contenuto scena
                if scene.content:
                    # Check if content is HTML (from rich text editor)
                    if scene.content.strip().startswith('<') or '<html>' in scene.content.lower():
                        # Parse HTML and create formatted runs
                        # Extract paragraphs from HTML
                        import re
                        # Split by <p> tags
                        p_tags = re.findall(r'<p[^>]*>(.*?)</p>', scene.content, re.DOTALL | re.IGNORECASE)

                        for p_content in p_tags:
                            if p_content.strip():
                                content_para = doc.add_paragraph()

                                # Use converter to add formatted runs
                                converter = HTMLToDocxConverter(content_para)
                                converter.feed(p_content)

                                self._format_paragraph(
                                    content_para,
                                    font_size=content_font_size,
                                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                                )
                    else:
                        # Plain text - preserve original logic
                        paragraphs = scene.content.split('\n\n')
                        for paragraph_text in paragraphs:
                            if paragraph_text.strip():
                                # Crea paragrafo e gestisci i newline come line breaks
                                content_para = doc.add_paragraph()

                                # Dividi per newline singoli (accapo)
                                lines = paragraph_text.split('\n')
                                for line_idx, line in enumerate(lines):
                                    if line.strip():
                                        run = content_para.add_run(line)
                                        # Aggiungi line break dopo ogni linea (tranne l'ultima)
                                        if line_idx < len(lines) - 1:
                                            run.add_break()

                                self._format_paragraph(
                                    content_para,
                                    font_size=content_font_size,
                                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                                )
                else:
                    # Placeholder se scena vuota
                    placeholder = doc.add_paragraph("[Scene content not yet written]")
                    self._format_paragraph(
                        placeholder,
                        font_size=content_font_size,
                        italic=True,
                        color=RGBColor(149, 165, 166)
                    )

            # Page break tra capitoli (se richiesto)
            if self._get_option('chapter_page_break', True) and i < len(chapters):
                doc.add_page_break()

        AppLogger.debug(f"Added manuscript content: {len(chapters)} chapters")

    def _format_paragraph(self, paragraph, font_size=None, bold=False, italic=False,
                         color=None, alignment=None):
        """
        Applica formattazione a un paragrafo

        Args:
            paragraph: Paragrafo da formattare
            font_size: Dimensione font in punti
            bold: Grassetto
            italic: Corsivo
            color: Colore RGB
            alignment: Allineamento
        """
        # Formatta tutti i run nel paragrafo
        for run in paragraph.runs:
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = bold
            if italic:
                run.font.italic = italic
            if color:
                run.font.color.rgb = color

        # Allineamento del paragrafo
        if alignment:
            paragraph.alignment = alignment
